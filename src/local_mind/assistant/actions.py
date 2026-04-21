from __future__ import annotations

import logging
import os
import platform
import re
import shutil
import subprocess
import sys
import time
import webbrowser
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote_plus

from local_mind.assistant.apps import AppIndex
from local_mind.config import settings

log = logging.getLogger(__name__)

# Known site aliases → URL. Matched against cleaned user queries.
SITE_ALIASES: dict[str, str] = {
    "gmail": "https://mail.google.com",
    "google mail": "https://mail.google.com",
    "google": "https://google.com",
    "youtube": "https://www.youtube.com",
    "yt": "https://www.youtube.com",
    "github": "https://github.com",
    "stackoverflow": "https://stackoverflow.com",
    "stack overflow": "https://stackoverflow.com",
    "twitter": "https://x.com",
    "x": "https://x.com",
    "reddit": "https://www.reddit.com",
    "netflix": "https://www.netflix.com",
    "amazon": "https://www.amazon.com",
    "linkedin": "https://www.linkedin.com",
    "whatsapp": "https://web.whatsapp.com",
    "slack": "https://slack.com",
    "discord": "https://discord.com/app",
    "notion": "https://www.notion.so",
    "chatgpt": "https://chat.openai.com",
    "claude": "https://claude.ai",
    "wikipedia": "https://en.wikipedia.org",
    "wiki": "https://en.wikipedia.org",
    "google maps": "https://maps.google.com",
    "maps": "https://maps.google.com",
    "google drive": "https://drive.google.com",
    "drive": "https://drive.google.com",
    "google docs": "https://docs.google.com",
    "docs.google": "https://docs.google.com",
    "google sheets": "https://sheets.google.com",
    "outlook": "https://outlook.live.com",
    "teams": "https://teams.microsoft.com",
    "zoom": "https://zoom.us",
    "spotify": "https://open.spotify.com",
    "news": "https://news.google.com",
    "google news": "https://news.google.com",
    "weather": "https://duckduckgo.com/?q=%21ducky+weather+today",
    "calendar": "https://calendar.google.com",
    "google calendar": "https://calendar.google.com",
}

# DuckDuckGo bang prefixes for site-scoped search.
SITE_BANGS: dict[str, str] = {
    "youtube": "yt",
    "yt": "yt",
    "github": "gh",
    "gh": "gh",
    "stackoverflow": "so",
    "so": "so",
    "reddit": "r",
    "twitter": "tw",
    "x": "tw",
    "wikipedia": "w",
    "wiki": "w",
    "amazon": "amazon",
    "maps": "gm",
    "google maps": "gm",
    "google": "g",
    "npm": "npm",
    "pypi": "pypi",
}

_URL_RE = re.compile(r"^https?://\S+$", re.IGNORECASE)
_DOMAIN_RE = re.compile(r"^(?:[a-z0-9-]+\.)+[a-z]{2,}(?:/\S*)?$", re.IGNORECASE)
_SAFE_NAME_RE = re.compile(r"[^A-Za-z0-9 _\-()]+")


@dataclass
class ActionResult:
    ok: bool
    speech: str
    detail: str = ""
    data: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "ok": self.ok,
            "speech": self.speech,
            "detail": self.detail,
            "data": self.data,
        }


class ActionRegistry:
    """Collection of assistant skills.

    Each skill returns an :class:`ActionResult`. Dangerous ones consult the
    ``confirm`` callback first (which can be a spoken "are you sure?" flow in
    the engine).
    """

    def __init__(self, apps: AppIndex | None = None) -> None:
        self.apps = apps or AppIndex()
        self.confirm_required = True

    # ── apps ────────────────────────────────────────────────────────────────
    def open_app(self, name: str) -> ActionResult:
        app, msg = self.apps.launch(name)
        return ActionResult(ok=bool(app), speech=msg, data={"app": app.name if app else None})

    def close_app(self, name: str) -> ActionResult:
        killed, msg = self.apps.close(name)
        return ActionResult(ok=bool(killed), speech=msg, data={"killed": killed})

    def list_running(self) -> ActionResult:
        procs = self.apps.list_running()
        top = sorted({p["name"] for p in procs})[:12]
        return ActionResult(
            ok=True,
            speech=f"You have {len(procs)} processes running. Top apps: {', '.join(top[:6])}.",
            data={"count": len(procs), "top": top},
        )

    # ── volume / media (via media keys) ─────────────────────────────────────
    def volume(self, direction: str, times: int = 2) -> ActionResult:
        try:
            from pynput.keyboard import Controller, Key
        except ImportError:
            return ActionResult(False, "pynput isn't installed, so I can't change the volume.")
        kb = Controller()
        direction = direction.lower()
        if direction in {"mute", "unmute"}:
            kb.press(Key.media_volume_mute)
            kb.release(Key.media_volume_mute)
            return ActionResult(True, "Toggled mute.")
        key = Key.media_volume_up if direction in {"up", "louder", "raise"} else Key.media_volume_down
        for _ in range(max(1, min(times, 20))):
            kb.press(key); kb.release(key); time.sleep(0.02)
        word = "up" if key == Key.media_volume_up else "down"
        return ActionResult(True, f"Volume {word}.")

    def media(self, command: str) -> ActionResult:
        try:
            from pynput.keyboard import Controller, Key
        except ImportError:
            return ActionResult(False, "pynput isn't installed.")
        kb = Controller()
        mapping = {
            "play": Key.media_play_pause,
            "pause": Key.media_play_pause,
            "toggle": Key.media_play_pause,
            "next": Key.media_next,
            "previous": Key.media_previous,
            "prev": Key.media_previous,
            "stop": getattr(Key, "media_stop", Key.media_play_pause),
        }
        key = mapping.get(command.lower())
        if key is None:
            return ActionResult(False, f"I don't know media command '{command}'.")
        kb.press(key); kb.release(key)
        return ActionResult(True, f"Media {command}.")

    # ── brightness ─────────────────────────────────────────────────────────
    def brightness(self, direction_or_value: str) -> ActionResult:
        try:
            import screen_brightness_control as sbc  # type: ignore
        except ImportError:
            return ActionResult(False, "screen-brightness-control isn't installed.")
        try:
            current_list = sbc.get_brightness()
            current = int(current_list[0]) if current_list else 50
        except Exception:
            current = 50
        val = direction_or_value.strip().lower()
        if val.isdigit() or (val.endswith("%") and val[:-1].isdigit()):
            target = int(val.rstrip("%"))
        elif val in {"up", "higher", "brighter"}:
            target = min(100, current + 15)
        elif val in {"down", "lower", "dimmer"}:
            target = max(0, current - 15)
        elif val in {"max", "full"}:
            target = 100
        elif val in {"min", "zero"}:
            target = 0
        else:
            return ActionResult(False, f"I don't understand brightness '{direction_or_value}'.")
        try:
            sbc.set_brightness(target)
        except Exception as e:
            return ActionResult(False, f"Couldn't change brightness: {e}")
        return ActionResult(True, f"Brightness set to {target} percent.", data={"percent": target})

    # ── info ───────────────────────────────────────────────────────────────
    def time_now(self) -> ActionResult:
        now = datetime.now()
        return ActionResult(True, f"It's {now.strftime('%I:%M %p').lstrip('0')}.", data={"iso": now.isoformat()})

    def date_now(self) -> ActionResult:
        now = datetime.now()
        return ActionResult(True, f"Today is {now.strftime('%A, %B %d, %Y')}.", data={"iso": now.isoformat()})

    def battery(self) -> ActionResult:
        try:
            import psutil
            b = psutil.sensors_battery()
        except Exception:
            b = None
        if b is None:
            return ActionResult(False, "Battery info isn't available on this machine.")
        pct = int(b.percent)
        plugged = b.power_plugged
        if plugged:
            return ActionResult(True, f"Battery at {pct} percent and plugged in.", data={"percent": pct, "plugged": True})
        return ActionResult(True, f"Battery at {pct} percent, on battery power.", data={"percent": pct, "plugged": False})

    def system_info(self) -> ActionResult:
        try:
            import psutil
            cpu = psutil.cpu_percent(interval=0.3)
            mem = psutil.virtual_memory()
        except Exception as e:
            return ActionResult(False, f"Can't read system info: {e}")
        gb_used = mem.used / (1024 ** 3)
        gb_total = mem.total / (1024 ** 3)
        text = (
            f"CPU at {cpu:.0f} percent. Memory {gb_used:.1f} of {gb_total:.1f} gigabytes. "
            f"Running {platform.system()} {platform.release()}."
        )
        return ActionResult(True, text, data={
            "cpu_percent": cpu,
            "mem_used_gb": round(gb_used, 2),
            "mem_total_gb": round(gb_total, 2),
            "platform": platform.platform(),
        })

    # ── open / navigate ─────────────────────────────────────────────────────
    def open_url(self, url: str) -> ActionResult:
        if not url:
            return ActionResult(False, "Open what URL?")
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
        try:
            webbrowser.open(url, new=2)
        except Exception as e:
            return ActionResult(False, f"Couldn't open the URL: {e}")
        return ActionResult(True, f"Opening {url}.", data={"url": url})

    def open_path(self, path: str) -> ActionResult:
        p = Path(path).expanduser()
        if not p.exists():
            return ActionResult(False, f"Path not found: {p}")
        try:
            _open_default(p)
        except Exception as e:
            return ActionResult(False, f"Couldn't open {p}: {e}")
        return ActionResult(True, f"Opening {p.name}.", data={"path": str(p)})

    def search_web(self, query: str) -> ActionResult:
        if not query:
            return ActionResult(False, "Search for what?")
        url = "https://duckduckgo.com/?q=" + quote_plus(query)
        webbrowser.open(url, new=2)
        return ActionResult(True, f"Searching the web for {query}.", data={"url": url, "query": query})

    def search_on_site(self, query: str, site: str) -> ActionResult:
        site_key = site.strip().lower()
        bang = SITE_BANGS.get(site_key)
        if not bang:
            for alias, b in SITE_BANGS.items():
                if alias in site_key:
                    bang = b
                    break
        if bang:
            url = f"https://duckduckgo.com/?q=%21{bang}+{quote_plus(query)}"
        else:
            url = f"https://duckduckgo.com/?q={quote_plus(query + ' ' + site)}"
        webbrowser.open(url, new=2)
        return ActionResult(True, f"Searching {site} for {query}.", data={"url": url, "site": site, "query": query})

    def open_target(self, query: str) -> ActionResult:
        """Unified 'open X' — handles URLs, domains, site aliases, then apps,
        falling back to a web search. This is what the wake phrase hits first."""
        raw = query.strip().rstrip(".?!")
        if not raw:
            return ActionResult(False, "Open what?")
        q = raw.lower()
        spoken_domain = re.sub(r"\s+dot\s+", ".", q)

        if _URL_RE.match(raw) or _URL_RE.match(spoken_domain):
            return self.open_url(raw if _URL_RE.match(raw) else spoken_domain)

        if _DOMAIN_RE.match(q) or _DOMAIN_RE.match(spoken_domain):
            return self.open_url(spoken_domain if _DOMAIN_RE.match(spoken_domain) else q)

        if q in SITE_ALIASES:
            return self.open_url(SITE_ALIASES[q])

        matches = self.apps.find(raw, top=1)
        if matches:
            app, msg = self.apps.launch(raw)
            return ActionResult(ok=bool(app), speech=msg, data={"app": app.name if app else None})

        for alias in sorted(SITE_ALIASES.keys(), key=lambda a: -len(a)):
            if alias == q or re.search(rf"\b{re.escape(alias)}\b", q):
                return self.open_url(SITE_ALIASES[alias])

        url = f"https://duckduckgo.com/?q=%21ducky+{quote_plus(raw)}"
        webbrowser.open(url, new=2)
        return ActionResult(True, f"I couldn't match '{raw}', so I searched the web.", data={"url": url, "query": raw})

    def refresh_apps(self) -> ActionResult:
        before = len(self.apps.refresh())
        apps = self.apps.force_refresh()
        delta = len(apps) - before
        if delta > 0:
            return ActionResult(True, f"Rescanned apps. Found {delta} new.", data={"total": len(apps), "new": delta})
        if delta < 0:
            return ActionResult(True, f"Rescanned apps. {-delta} removed.", data={"total": len(apps), "removed": -delta})
        return ActionResult(True, f"Rescanned apps. {len(apps)} known.", data={"total": len(apps)})

    def learn_url(self, url: str) -> ActionResult:
        if not url:
            return ActionResult(False, "Learn from what URL?")
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
        try:
            from local_mind.knowledge import knowledge
            res = knowledge.learn_url(url)
        except Exception as e:
            return ActionResult(False, f"Couldn't learn from {url}: {e}")
        if res.get("status") == "learned":
            return ActionResult(True, f"Got it, learned {res.get('chunks', 0)} chunks from that page.", data=res)
        return ActionResult(False, f"Skipped: {res.get('reason', 'unknown')}", data=res)

    # ── typing / keys ──────────────────────────────────────────────────────
    def type_text(self, text: str) -> ActionResult:
        if not text:
            return ActionResult(False, "Type what?")
        try:
            from pynput.keyboard import Controller
        except ImportError:
            return ActionResult(False, "pynput isn't installed.")
        Controller().type(text)
        return ActionResult(True, f"Typed {len(text)} characters.", data={"length": len(text)})

    def press_keys(self, combo: str) -> ActionResult:
        try:
            from pynput.keyboard import Controller, Key
        except ImportError:
            return ActionResult(False, "pynput isn't installed.")
        parts = [p.strip().lower() for p in combo.replace(" ", "").split("+") if p.strip()]
        if not parts:
            return ActionResult(False, "Which keys?")
        kb = Controller()
        keys: list[Any] = []
        for p in parts:
            if hasattr(Key, p):
                keys.append(getattr(Key, p))
            elif len(p) == 1:
                keys.append(p)
            else:
                return ActionResult(False, f"Unknown key '{p}'.")
        for k in keys:
            kb.press(k)
        for k in reversed(keys):
            kb.release(k)
        return ActionResult(True, f"Pressed {combo}.", data={"combo": combo})

    def screenshot(self, save_dir: Path | None = None) -> ActionResult:
        try:
            from PIL import ImageGrab  # type: ignore
        except ImportError:
            return ActionResult(False, "Pillow isn't installed, so I can't take screenshots.")
        save_dir = Path(save_dir or (Path.home() / "Pictures" / "LocalMind"))
        save_dir.mkdir(parents=True, exist_ok=True)
        fname = save_dir / f"screenshot-{datetime.now().strftime('%Y%m%d-%H%M%S')}.png"
        try:
            img = ImageGrab.grab()
            img.save(fname)
        except Exception as e:
            return ActionResult(False, f"Screenshot failed: {e}")
        return ActionResult(True, f"Saved screenshot to {fname.name}.", data={"path": str(fname)})

    # ── session control (dangerous) ────────────────────────────────────────
    def lock(self) -> ActionResult:
        if sys.platform != "win32":
            return ActionResult(False, "Lock is only implemented for Windows right now.")
        try:
            subprocess.run(["rundll32.exe", "user32.dll,LockWorkStation"], check=False)
        except Exception as e:
            return ActionResult(False, f"Couldn't lock: {e}")
        return ActionResult(True, "Locking the workstation.")

    def sleep_system(self) -> ActionResult:
        if sys.platform != "win32":
            return ActionResult(False, "Sleep is only implemented for Windows right now.")
        try:
            subprocess.Popen(["rundll32.exe", "powrprof.dll,SetSuspendState", "0,1,0"])
        except Exception as e:
            return ActionResult(False, f"Couldn't sleep: {e}")
        return ActionResult(True, "Putting the system to sleep.")

    def shutdown(self, delay: int = 10) -> ActionResult:
        if sys.platform == "win32":
            cmd = ["shutdown", "/s", "/t", str(delay)]
        elif sys.platform == "darwin":
            cmd = ["osascript", "-e", 'tell app "System Events" to shut down']
        else:
            cmd = ["shutdown", "-h", f"+{max(1, delay // 60)}"]
        try:
            subprocess.Popen(cmd)
        except Exception as e:
            return ActionResult(False, f"Couldn't shut down: {e}")
        return ActionResult(True, f"Shutting down in {delay} seconds. Say 'cancel shutdown' to abort.")

    def cancel_shutdown(self) -> ActionResult:
        if sys.platform != "win32":
            return ActionResult(False, "Not supported on this platform.")
        try:
            subprocess.run(["shutdown", "/a"], check=False)
        except Exception as e:
            return ActionResult(False, f"Couldn't cancel: {e}")
        return ActionResult(True, "Shutdown cancelled.")

    def reboot(self, delay: int = 10) -> ActionResult:
        if sys.platform == "win32":
            cmd = ["shutdown", "/r", "/t", str(delay)]
        else:
            cmd = ["shutdown", "-r", "now"]
        try:
            subprocess.Popen(cmd)
        except Exception as e:
            return ActionResult(False, f"Couldn't reboot: {e}")
        return ActionResult(True, f"Rebooting in {delay} seconds.")

    # ── documents ───────────────────────────────────────────────────────────
    def create_document(
        self,
        name: str | None,
        fmt: str = "auto",
        content: str | None = None,
        open_after: bool = True,
    ) -> ActionResult:
        """Create an empty (or templated) document and optionally open it.

        Supports formats: 'txt', 'md', 'docx', 'pdf'. Passing 'auto' picks
        .docx if python-docx is importable, otherwise .txt.
        """
        chosen = _resolve_format(fmt)
        title = (name or "Untitled").strip() or "Untitled"
        safe_name = _safe_filename(title) or "Untitled"
        docs_dir = settings.assistant_docs_dir
        docs_dir.mkdir(parents=True, exist_ok=True)
        path = _unique_path(docs_dir / f"{safe_name}.{chosen}")

        body = content or ""
        try:
            if chosen == "txt":
                path.write_text(body, encoding="utf-8")
            elif chosen == "md":
                header = f"# {title}\n\n"
                path.write_text(header + body, encoding="utf-8")
            elif chosen == "docx":
                if not _have("docx"):
                    return ActionResult(False, "python-docx isn't installed, so I can't create Word documents.")
                _write_docx(path, title, body)
            elif chosen == "pdf":
                if not _have("reportlab"):
                    return ActionResult(False, "reportlab isn't installed, so I can't create PDFs.")
                _write_pdf(path, title, body)
            else:
                return ActionResult(False, f"I don't know how to create a .{chosen} file.")
        except Exception as e:
            log.exception("Document creation failed")
            return ActionResult(False, f"Couldn't create the document: {e}")

        if open_after:
            try:
                _open_default(path)
            except Exception as e:
                log.debug("open after create failed: %s", e)

        return ActionResult(
            True,
            f"Created {path.name} in {docs_dir.name}.",
            detail=str(path),
            data={"path": str(path), "format": chosen, "title": title},
        )

    def write_document(
        self,
        topic: str,
        fmt: str = "auto",
        name: str | None = None,
        open_after: bool = True,
    ) -> ActionResult:
        """Have the loaded LLM draft a document on ``topic`` and save it."""
        topic = (topic or "").strip()
        if not topic:
            return ActionResult(False, "What should I write about?")
        try:
            from local_mind.chat import smart_chat
            from local_mind.models import model_manager
        except Exception as e:
            return ActionResult(False, f"LLM isn't available: {e}")
        if not model_manager.loaded:
            return ActionResult(False, "I need a language model loaded. Open the chat tab to load one.")

        prompt = (
            f"Write a well-structured, comprehensive document about: {topic}\n\n"
            "Format your response in markdown:\n"
            "- Start with a short one-line title (no leading #).\n"
            "- Use '## Heading' for each major section.\n"
            "- Use '- ' for bullet points and blank lines between paragraphs.\n"
            "- Be factual, clear, and thorough.\n"
        )

        try:
            result = smart_chat(prompt, history=None, use_rag=True, temperature=0.5, max_tokens=1600)
        except Exception as e:
            return ActionResult(False, f"The model errored out: {e}")
        content = (result.get("content") or "").strip()
        if not content:
            return ActionResult(False, "The model didn't return anything.")

        title = name
        if not title:
            first_line, _, rest = content.partition("\n")
            first_line = first_line.strip().lstrip("# ").strip()
            if 2 <= len(first_line) <= 80:
                title = first_line
                content = rest.lstrip()
        if not title:
            title = topic[:60].rstrip(" .,:;")

        created = self.create_document(name=title, fmt=fmt, content=content, open_after=open_after)
        if created.ok:
            created.speech = f"I wrote a {created.data.get('format', fmt)} about {topic}. Saved as {Path(created.data['path']).name}."
        return created

    # ── llm fallback ────────────────────────────────────────────────────────
    def ask_llm(self, question: str) -> ActionResult:
        try:
            from local_mind.chat import smart_chat
            from local_mind.models import model_manager
        except Exception as e:
            return ActionResult(False, f"LLM isn't available: {e}")
        if not model_manager.loaded:
            return ActionResult(False, "I don't have a language model loaded. Use the chat tab to load one.")
        try:
            result = smart_chat(question, history=None, use_rag=True, temperature=0.4, max_tokens=320)
        except Exception as e:
            return ActionResult(False, f"The model errored out: {e}")
        answer = (result.get("content") or "").strip()
        if not answer:
            return ActionResult(False, "I didn't get an answer from the model.")
        return ActionResult(True, _shorten_for_speech(answer), detail=answer, data=result)


def _shorten_for_speech(text: str, max_chars: int = 420) -> str:
    text = " ".join(text.split())
    if len(text) <= max_chars:
        return text
    head = text[:max_chars]
    for sep in (". ", "! ", "? "):
        idx = head.rfind(sep)
        if idx > 120:
            return head[: idx + 1]
    return head.rstrip() + "…"


# ── document helpers ────────────────────────────────────────────────────────
def _have(module: str) -> bool:
    try:
        if module == "docx":
            import docx  # noqa: F401
        elif module == "reportlab":
            import reportlab  # noqa: F401
        else:
            return False
        return True
    except Exception:
        return False


def _resolve_format(fmt: str) -> str:
    fmt = (fmt or "auto").strip().lower().lstrip(".")
    if fmt == "auto":
        return "docx" if _have("docx") else "txt"
    aliases = {
        "text": "txt", "plaintext": "txt", "plain": "txt", "note": "txt",
        "markdown": "md", "md": "md",
        "word": "docx", "word document": "docx", "doc": "docx", "docx": "docx",
        "pdf": "pdf",
    }
    return aliases.get(fmt, fmt)


def _safe_filename(name: str) -> str:
    cleaned = _SAFE_NAME_RE.sub(" ", name).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned[:80]


def _unique_path(base: Path) -> Path:
    if not base.exists():
        return base
    stem, ext, parent = base.stem, base.suffix, base.parent
    for i in range(2, 100):
        cand = parent / f"{stem} ({i}){ext}"
        if not cand.exists():
            return cand
    return parent / f"{stem}-{int(time.time())}{ext}"


def _iter_blocks(body: str):
    for block in re.split(r"\n\s*\n", body):
        b = block.strip()
        if b:
            yield b


def _write_docx(path: Path, title: str, body: str) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading(title, level=0)
    for block in _iter_blocks(body or ""):
        if block.startswith("### "):
            doc.add_heading(block[4:].strip(), level=3)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=1)
        elif all(line.lstrip().startswith(("- ", "* ")) for line in block.splitlines()):
            for line in block.splitlines():
                doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(block)
    if not body.strip():
        doc.add_paragraph()
    doc.save(str(path))


def _write_pdf(path: Path, title: str, body: str) -> None:
    from reportlab.lib.pagesizes import letter  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.lib.units import inch  # type: ignore
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore

    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(_xml_escape(title), styles["Title"]), Spacer(1, 0.18 * inch)]
    for block in _iter_blocks(body or ""):
        if block.startswith("### "):
            story.append(Paragraph(_xml_escape(block[4:].strip()), styles["Heading3"]))
        elif block.startswith("## "):
            story.append(Paragraph(_xml_escape(block[3:].strip()), styles["Heading2"]))
        elif block.startswith("# "):
            story.append(Paragraph(_xml_escape(block[2:].strip()), styles["Heading1"]))
        elif all(line.lstrip().startswith(("- ", "* ")) for line in block.splitlines()):
            for line in block.splitlines():
                story.append(Paragraph(f"• {_xml_escape(line.lstrip()[2:].strip())}", styles["BodyText"]))
        else:
            story.append(Paragraph(_xml_escape(block).replace("\n", "<br/>"), styles["BodyText"]))
            story.append(Spacer(1, 0.08 * inch))
    if len(story) == 2:
        story.append(Paragraph(" ", styles["BodyText"]))
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        title=title,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )
    doc.build(story)


def _xml_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _open_default(path: Path) -> None:
    if sys.platform == "win32":
        os.startfile(str(path))  # type: ignore[attr-defined]
    elif sys.platform == "darwin":
        subprocess.Popen(["open", str(path)])
    else:
        subprocess.Popen(["xdg-open", str(path)])
